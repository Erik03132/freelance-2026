#!/usr/bin/env python3
"""Convert markdown file to DOC format."""

import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_md_to_doc(md_file, doc_file):
    """Convert markdown file to DOC format."""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse markdown to HTML
    html = markdown.markdown(md_content, extensions=['tables'])
    
    # Process content line by line for better control
    lines = md_content.split('\n')
    current_list = None
    table_rows = []
    in_table = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            if in_table and table_rows:
                # Add table
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue
        
        # Horizontal rule
        if stripped.startswith('***') or stripped.startswith('---'):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        # Headers
        if stripped.startswith('# '):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = stripped[2:].replace('📅', '').replace('📋', '').replace('📞', '').strip()
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        if stripped.startswith('## '):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = stripped[3:].replace('📅', '').replace('📋', '').replace('📞', '').strip()
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        if stripped.startswith('### '):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = stripped[4:].replace('📅', '').replace('📋', '').replace('📞', '').strip()
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Table detection
        if '|' in stripped and stripped.count('|') >= 2:
            in_table = True
            # Check if next line is separator
            if i + 1 < len(lines) and '|' in lines[i + 1] and ':' in lines[i + 1]:
                table_rows.append([cell.strip() for cell in stripped.split('|')[1:-1]])
                i += 2  # Skip header and separator
                # Read table rows
                while i < len(lines) and '|' in lines[i] and ':' not in lines[i]:
                    table_rows.append([cell.strip() for cell in lines[i].split('|')[1:-1]])
                    i += 1
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
                continue
            else:
                table_rows.append([cell.strip() for cell in stripped.split('|')[1:-1]])
                i += 1
                continue
        
        # Bullet lists
        if stripped.startswith('- '):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            if current_list != 'bullet':
                current_list = 'bullet'
            text = stripped[2:].replace('**', '').replace('`', '')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            i += 1
            continue
        
        # Checkbox lists
        if stripped.startswith('- [ ]'):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = stripped[5:].replace('**', '').replace('`', '')
            p = doc.add_paragraph()
            run = p.add_run('☐ ' + text)
            run.font.size = Pt(10)
            i += 1
            continue
        
        # Code blocks
        if stripped.startswith('```'):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            # Skip code block
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue
        
        # Blockquotes
        if stripped.startswith('>'):
            if in_table and table_rows:
                add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = stripped[1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue
        
        # Regular paragraphs
        if in_table and table_rows:
            add_table_to_doc(doc, table_rows)
            table_rows = []
            in_table = False
        
        current_list = None
        text = stripped.replace('**', '').replace('`', '')
        # Remove emoji
        text = re.sub(r'[📅📋📞]', '', text)
        p = doc.add_paragraph()
        p.add_run(text)
        p.paragraph_format.space_after = Pt(6)
        i += 1
    
    # Add any remaining table
    if table_rows:
        add_table_to_doc(doc, table_rows)
    
    # Save document
    doc.save(doc_file)
    print(f"Successfully converted {md_file} to {doc_file}")

def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:  # Header row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

if __name__ == '__main__':
    convert_md_to_doc('5-шагов-создания-ООО-ветеранов.md', '5-шагов-создания-ООО-ветеранов.doc')
